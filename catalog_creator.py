import tkinter as tk
from tkinter import messagebox, filedialog, ttk, simpledialog
import time
import json
import os
import pandas as pd
from docx import Document

class MyApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Catalog Creator")
        self.root.geometry("400x400")

        
        self.center_window()

        
        self.loader_label = tk.Label(root, text="Welcome to beta version", font=("Arial", 16))
        self.loader_label.pack(pady=20)

        
        self.fade_in()

        
        self.create_button = self.create_button_with_hover("Создать каталог", self.open_catalog_window)
        self.create_button.pack(pady=10)

        self.save_button = self.create_button_with_hover("Сохранить", self.save_file)
        self.save_button.pack(pady=10)

        self.view_button = self.create_button_with_hover("Посмотреть последние каталоги", self.view_recent_catalogs)
        self.view_button.pack(pady=10)

        self.import_button = self.create_button_with_hover("Импортировать каталог", self.import_catalog)
        self.import_button.pack(pady=10)

        self.export_button = self.create_button_with_hover("Экспортировать каталог", self.export_catalog)
        self.export_button.pack(pady=10)

        self.exit_button = self.create_button_with_hover("Выход", self.exit_app)
        self.exit_button.pack(pady=10)

     
        self.info_button = self.create_info_button()
        self.info_button.place(relx=1.0, rely=1.0, anchor='se', x=-10, y=-10)

        
        self.support_button = self.create_support_button()
        self.support_button.place(relx=0.0, rely=1.0, anchor='sw', x=10, y=-10)

        
        self.movie_data = []
        self.book_data = []
        self.animation_data = []

        # Хранилище для вкладок
        self.notebook = None

    def center_window(self):
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - 400) // 2
        y = (screen_height - 300) // 2
        self.root.geometry(f"+{x}+{y}")

    def fade_in(self, step=0.1, delay=0.05):
        for i in range(1, 11):
            self.root.attributes("-alpha", i * step)
            self.root.update()
            time.sleep(delay)

    def create_button_with_hover(self, text, command):
        """Создание кнопки с эффектом наведения."""
        button = tk.Button(self.root, text=text, command=command)

        def on_enter(e):
            button['background'] = 'lightblue'

        def on_leave(e):
            button['background'] = 'SystemButtonFace'

        button.bind("<Enter>", on_enter)
        button.bind("<Leave>", on_leave)

        return button

    def open_catalog_window(self):
        catalog_window = tk.Toplevel(self.root)
        catalog_window.title("Catalog Window")
        catalog_window.geometry("600x400")

        # Создание вкладок
        self.notebook = ttk.Notebook(catalog_window)

        # Изначальные вкладки
        self.create_new_tab(tab_name="Фильмы", data=self.movie_data)
        self.create_new_tab(tab_name="Книги", data=self.book_data)
        self.create_new_tab(tab_name="Анимация", data=self.animation_data)

        self.notebook.pack(fill='both', expand=True)

        # Кнопка для создания новой вкладки
        new_tab_button = tk.Button(catalog_window, text="+", command=self.create_new_tab_dialog)
        new_tab_button.place(relx=1.0, rely=0.0, anchor='ne', x=-10, y=10)  # Правый верхний угол

        # Кнопка для удаления текущей вкладки
        delete_tab_button = tk.Button(catalog_window, text="-", command=self.delete_current_tab)
        delete_tab_button.place(relx=1.0, rely=0.0, anchor='ne', x=-50, y=10)  # Правый верхний угол, смещенная влево

    def create_new_tab_dialog(self):
        """Создает новую вкладку по запросу пользователя."""
        tab_name = simpledialog.askstring("Название вкладки", "Введите название новой вкладки:")
        if tab_name:
            self.create_new_tab(tab_name=tab_name)

    def create_new_tab(self, tab_name="Новая вкладка", data=None):
        """Создает новую вкладку в ноутбуке."""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text=tab_name)

        
        treeview = ttk.Treeview(tab, columns=("name", "genre", "year"))
        self.setup_treeview(treeview, ["Название", "Жанр", "Год"])
        self.populate_treeview(treeview, data or [])
        treeview.pack(fill='both', expand=True)

        
        add_button = tk.Button(tab, text="Добавить пункт", command=lambda: self.add_item(treeview, tab_name))
        add_button.pack(pady=5)

        # Добавление событий двойного клика для редактирования
        treeview.bind("<Double-1>", lambda event: self.edit_item(treeview, tab_name))

    def setup_treeview(self, treeview, columns):
        for i, col in enumerate(columns):
            treeview.column(i, width=150, anchor='center')
            treeview.heading(i, text=col)

        treeview.column('#0', width=0, stretch=tk.NO)
        treeview['show'] = 'headings'

    def populate_treeview(self, treeview, data):
        treeview.delete(*treeview.get_children())  # Очистка текущих данных
        for item in data:
            values = tuple(item.values())
            treeview.insert('', 'end', values=values)

    def add_item(self, treeview, tab_name):
        """Добавляет новый пункт в список."""
        name = simpledialog.askstring("Введите название", "Название:")
        genre = simpledialog.askstring("Введите жанр", "Жанр:")
        year = simpledialog.askstring("Введите год", "Год:")

        if name and genre and year:
            new_item = {"name": name, "genre": genre, "year": year}
            if tab_name == "Фильмы":
                self.movie_data.append(new_item)
            elif tab_name == "Книги":
                self.book_data.append(new_item)
            elif tab_name == "Анимация":
                self.animation_data.append(new_item)
            self.populate_treeview(treeview, self.get_current_data(tab_name))

    def get_current_data(self, tab_name):
        """Возвращает данные для текущей вкладки."""
        if tab_name == "Фильмы":
            return self.movie_data
        elif tab_name == "Книги":
            return self.book_data
        elif tab_name == "Анимация":
            return self.animation_data
        return []

    def edit_item(self, treeview, tab_name):
        selected_item = treeview.selection()
        if selected_item:
            item_values = treeview.item(selected_item, 'values')

            # Создание нового окна для редактирования
            edit_window = tk.Toplevel(self.root)
            edit_window.title("Редактировать элемент")

            fields = [("Название", item_values[0]), ("Жанр", item_values[1]), ("Год", item_values[2])]
            entries = []

            for field, value in fields:
                label = tk.Label(edit_window, text=field)
                label.pack(pady=5)
                entry = tk.Entry(edit_window)
                entry.insert(0, value)
                entry.pack(pady=5)
                entries.append(entry)

            def save_changes():
                for i, entry in enumerate(entries):
                    # Обновление данных в соответствующем списке
                    if tab_name == "Фильмы":
                        self.movie_data[treeview.index(selected_item)][list(self.movie_data[0].keys())[i]] = entry.get()
                    elif tab_name == "Книги":
                        self.book_data[treeview.index(selected_item)][list(self.book_data[0].keys())[i]] = entry.get()
                    elif tab_name == "Анимация":
                        self.animation_data[treeview.index(selected_item)][list(self.animation_data[0].keys())[i]] = entry.get()
                self.populate_treeview(treeview, self.get_current_data(tab_name))
                edit_window.destroy()

            save_button = tk.Button(edit_window, text="Сохранить", command=save_changes)
            save_button.pack(pady=10)

    def delete_current_tab(self):
        """Удаляет текущую вкладку."""
        selected_tab_index = self.notebook.index(self.notebook.select())
        if selected_tab_index >= 0:
            self.notebook.forget(selected_tab_index)
        else:
            messagebox.showwarning("Предупреждение", "Сначала выберите вкладку для удаления.")

    def save_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".json",
                                                   filetypes=[("JSON files", "*.json"), ("All files", "*.*")])
        if file_path:
            data_to_save = {
                "movies": self.movie_data,
                "books": self.book_data,
                "animations": self.animation_data
            }
            with open(file_path, 'w') as f:
                json.dump(data_to_save, f)
            messagebox.showinfo("Сохранение", "Данные успешно сохранены!")

    def view_recent_catalogs(self):
        # Здесь можно реализовать просмотр последних каталогов
        messagebox.showinfo("Посмотреть каталоги", "Функция просмотра недоступна в этой версии.")

    def import_catalog(self):
        file_path = filedialog.askopenfilename(filetypes=[
            ("JSON files", "*.json"),
            ("Word files", "*.docx"),
            ("Excel files", "*.xlsx"),
            ("All files", "*.*")])
        if file_path:
            if file_path.endswith('.json'):
                with open(file_path, 'r') as f:
                    data = json.load(f)
                    self.movie_data = data.get("movies", [])
                    self.book_data = data.get("books", [])
                    self.animation_data = data.get("animations", [])
                    messagebox.showinfo("Импорт", "Данные успешно импортированы из JSON!")
            elif file_path.endswith('.docx'):
                self.import_from_word(file_path)
            elif file_path.endswith('.xlsx'):
                self.import_from_excel(file_path)

    def import_from_word(self, file_path):
        """Импортирует данные из документа Word."""
        doc = Document(file_path)
        data = {"movies": [], "books": [], "animations": []}
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith("Фильмы:"):
                category = "movies"
            elif paragraph.text.startswith("Книги:"):
                category = "books"
            elif paragraph.text.startswith("Анимация:"):
                category = "animations"
            else:
                if category:
                    parts = paragraph.text.split(",")
                    if len(parts) == 3:
                        item = {"name": parts[0].strip(), "genre": parts[1].strip(), "year": parts[2].strip()}
                        data[category].append(item)

        self.movie_data = data.get("movies", [])
        self.book_data = data.get("books", [])
        self.animation_data = data.get("animations", [])
        messagebox.showinfo("Импорт", "Данные успешно импортированы из Word!")

    def import_from_excel(self, file_path):
        """Импортирует данные из Excel файла."""
        xls = pd.ExcelFile(file_path)
        data = {"movies": [], "books": [], "animations": []}
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name)
            for index, row in df.iterrows():
                item = {"name": row['Название'], "genre": row['Жанр'], "year": row['Год']}
                if sheet_name == "Фильмы":
                    data['movies'].append(item)
                elif sheet_name == "Книги":
                    data['books'].append(item)
                elif sheet_name == "Анимация":
                    data['animations'].append(item)

        self.movie_data = data.get("movies", [])
        self.book_data = data.get("books", [])
        self.animation_data = data.get("animations", [])
        messagebox.showinfo("Импорт", "Данные успешно импортированы из Excel!")

    def export_catalog(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".json",
                                                   filetypes=[("JSON files", "*.json"), ("Word files", "*.docx"), ("Excel files", "*.xlsx"), ("All files", "*.*")])
        if file_path:
            if file_path.endswith('.json'):
                self.export_to_json(file_path)
            elif file_path.endswith('.docx'):
                self.export_to_word(file_path)
            elif file_path.endswith('.xlsx'):
                self.export_to_excel(file_path)

    def export_to_json(self, file_path):
        """Экспорт данных в формате JSON."""
        data_to_save = {
            "movies": self.movie_data,
            "books": self.book_data,
            "animations": self.animation_data
        }
        with open(file_path, 'w') as f:
            json.dump(data_to_save, f)
        messagebox.showinfo("Экспорт", "Данные успешно экспортированы в JSON!")

    def export_to_word(self, file_path):
        """Экспорт данных в документ Word."""
        doc = Document()
        doc.add_heading('Каталог', 0)

        if self.movie_data:
            doc.add_heading('Фильмы:', level=1)
            for item in self.movie_data:
                doc.add_paragraph(f"{item['name']}, {item['genre']}, {item['year']}")

        if self.book_data:
            doc.add_heading('Книги:', level=1)
            for item in self.book_data:
                doc.add_paragraph(f"{item['name']}, {item['genre']}, {item['year']}")

        if self.animation_data:
            doc.add_heading('Анимация:', level=1)
            for item in self.animation_data:
                doc.add_paragraph(f"{item['name']}, {item['genre']}, {item['year']}")

        doc.save(file_path)
        messagebox.showinfo("Экспорт", "Данные успешно экспортированы в Word!")

    def export_to_excel(self, file_path):
        """Экспорт данных в Excel файл."""
        with pd.ExcelWriter(file_path) as writer:
            if self.movie_data:
                df_movies = pd.DataFrame(self.movie_data)
                df_movies.to_excel(writer, sheet_name='Фильмы', index=False)

            if self.book_data:
                df_books = pd.DataFrame(self.book_data)
                df_books.to_excel(writer, sheet_name='Книги', index=False)

            if self.animation_data:
                df_animations = pd.DataFrame(self.animation_data)
                df_animations.to_excel(writer, sheet_name='Анимация', index=False)

        messagebox.showinfo("Экспорт", "Данные успешно экспортированы в Excel!")

    def exit_app(self):
        self.root.quit()

    def create_info_button(self):
        button = tk.Button(self.root, text="i", command=self.show_info)
        button.configure(bg="lightgrey", fg="black", font=("Arial", 10, "bold"))
        return button

    def show_info(self):
        messagebox.showinfo("Информация", "@whyshelton")

    def create_support_button(self):
        button = tk.Button(self.root, text="Поддержка", command=self.show_support)
        button.configure(bg="lightgrey", fg="black", font=("Arial", 10, "bold"))
        return button

    def show_support(self):
        messagebox.showinfo("Поддержка", "@whyshelton")

if __name__ == "__main__":
    root = tk.Tk()
    app = MyApp(root)
    root.mainloop()
